#!/usr/bin/env python
# coding: utf-8

# Matthew Client Management System
# By Matthew Stubenberg
# Purpose is to create a CMS to store information about clients and generate word documents.


import csv
from os.path import exists
import docx
import os

def createWelcomeLetter(person):
    #This function is only one line but in theory we could add more things specifically related to this letter in the future.
    generateDocument(person,'Client Letter',person['client_id'])

def generateDocument(template_values,docname,clientid):
    #This is a function that replaces values in a word document.
    #It takes a dictionary and replaces all of the keys (with [ and ] added to the ends) and replaces them with the values
    #This way we don't have to maually type in every key we want to replace.
    #It also pulls and saves documents to and from specific locations to keep things organized.
    doc = docx.Document("Templates/" + docname + ".docx")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    paragraphs = doc.paragraphs

    for paragraph in paragraphs:
        for key,value in template_values.items():
            if(key in paragraph.text):
                paragraph.text = paragraph.text.replace('[' + key + ']',value)
    filename = 'Clients/' + str(clientid) + "/" + docname + ".docx"
    doc.save(filename)
    print("Created",docname)
def addNewClient(fieldnames):
    #This function asks the user for client information and calls the addClient function to actually update the csv
    person = {}
    person['client_id'] = str(getNumberOfCurrentClients())
    for item in fieldnames:
        #We want to add the Client ID ourselves so we want to skip it if it's asked.
        if(item == 'client_id'):
            continue
        while True:
            #We use a while loop to make sure we keep asking until the user inputs something for every field.
            userinput = input("Enter " + item + ":")
            if(userinput.strip() != ""):
                #This checks to make sure the user types in something for every field.
                break
            else:
                print("You must enter a value.")
        person[item] = userinput

    addClient(person)
    createWelcomeLetter(person)
    print("New client added successfully")

def getNumberOfCurrentClients():
    #Opens the clients csv and gets the length of the returned list
    with open('clients.csv') as csvfile:
        reader = csv.reader(csvfile, delimiter=' ', quotechar='|')
        return len(list(reader))
def createClientsCSV(fieldnames):
    #Creates a new csv if it doesn't exist
    #This way we don't hit file not found issues
    #This is run at the beginning of every run
    if(exists('clients.csv') == False):
        with open('clients.csv', 'w', newline='') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
    if(os.path.isdir('Clients') == False):
        os.mkdir('Clients') #Makes the clients folder if it doesn't already exist
def addClient(person):
    #Appends to the bottom of the csv a new client
    with open('clients.csv', 'a', newline='') as csvfile:
        fieldnames = person.keys()
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writerow(person)

    #Make a directory to hold that clients information
    if (os.path.isdir('Clients/' + person['client_id']) == False):
        os.mkdir('Clients/' + person['client_id'])
def viewClient():
    with open('clients.csv', newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        clients = list(reader)
    #Print all of the clients and ask the user which person they want to edit
    for client in clients:
        print(client['client_id'],":",client['first_name'],client['last_name'])
    while True:
        client_id = input("Enter the number of the client you would like to view: ")
        if(client_id.strip() == ""):
            #Make sure it's not just an empty string
            print("Please enter a value.")
        elif(client_id.isdigit() == False):
            #Make sure it's a number
            print("Please enter a number.")
        elif(int(client_id) < 1 or int(client_id) > len(clients)):
            #Make sure the number falls in the range of allowable client ids.
            print("Number not recognized")
        else:
            break

    #Save the individual client dictionary as person
    #We have to subtract 1 since the first client has an ID of 1 but an index of 0.
    person = clients[int(client_id) - 1]
    for key,value in client.items():
        print(key,":",value)
def updateClient():
    #Asks the user what item they would like to update
    #Updates the dictionary of that item
    #Overwrites the entire csv with the pull clients list.
    #This would not scale very effectively
    print("Which client would you like to update?")

    #Open the client csv, create a list of dictionary objects for the clients
    with open('clients.csv', newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        clients = list(reader)
    #Print all of the clients and ask the user which person they want to edit
    for client in clients:
        print(client['client_id'],":",client['first_name'],client['last_name'])
    while True:
        client_id = input("Enter the number of the client you would like to update: ")
        if(client_id.strip() == ""):
            #Make sure it's not just an empty string
            print("Please enter a value.")
        elif(client_id.isdigit() == False):
            #Make sure it's a number
            print("Please enter a number.")
        elif(int(client_id) < 1 or int(client_id) > len(clients)):
            #Make sure the number falls in the range of allowable client ids.
            print("Number not recognized")
        else:
            break

    #Save the individual client dictionary as person
    #We have to subtract 1 since the first client has an ID of 1 but an index of 0.
    person = clients[int(client_id) - 1]

    #Loop through the keys of the person and ask the user what they want to update.
    print("What would you like to update?")
    x=0
    personkeys = []
    for key,value in client.items():
        print(x+1,":",key,":",value)
        personkeys.append(key)
        x+=1
    itemtoupdate = input("Enter number of item you want to update: ")

    #Get updated value and update the that key in the person dictionary with that new value
    newitem = input("What do you want to change " + personkeys[int(itemtoupdate) - 1] + " to?: ")
    person[personkeys[int(itemtoupdate) - 1]] = newitem

    #Replace that item in the clients list with the updated person dictionary
    clients[int(person['client_id']) - 1] = person

    #Overwrite the clients csv and rewrite all the clients that we pulled in the beginning of this function.
    #This would not scale. If you had 10,000 clients you'd have to write that dictionary on each update.
    with open('clients.csv', 'w', newline='') as csvfile:
        fieldnames = clients[0].keys()
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for client in clients:
            writer.writerow(client)
    print("Client",person['first_name'],person['last_name'],"has been udpated")

#Start of Program
client_csv_fieldnames = ['client_id','first_name', 'last_name','dob','address','city','state','zip']
createClientsCSV(client_csv_fieldnames) #Creates a client csv if one doesn't already exist
print("Welcome to Matthew CMS")
#While loop allows the user to continue to use the program until they are done
while True:
    print("") #Adds a blank line to better seperate the text
    print("What do you want to do?")
    print("1: View Client")
    print("2: Add New Client")
    print("3: Update Client")
    print("4: Exit")
    option = input("Enter Option: ")
    #The selected option calls the various function responsible for that feature.
    if(option == '1'):
        viewClient()
    elif(option == '2'):
        addNewClient(client_csv_fieldnames)
    elif(option == '3'):
        updateClient()
    elif(option == '4'):
        break
    else:
        print("Sorry we didn't understand your input.")

print("Thank you for using the Matthew CMS")

